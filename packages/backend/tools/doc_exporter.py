from __future__ import annotations
import re
from pathlib import Path
import docx
from docx.shared import Inches, Pt

def _add_markdown_paragraph(parent, line: str, style: str | None = None):
    """Add a paragraph to a document or cell, parsing inline markdown styling (**bold**, *italic*, `code`)."""
    # Create the paragraph
    if hasattr(parent, 'add_paragraph'):
        p = parent.add_paragraph(style=style)
    else:
        p = parent.paragraphs[0]
        if style:
            p.style = style

    # Regular expression to tokenize bold, italic, inline code, and text
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|[^_*`]+)')
    tokens = pattern.findall(line)
    
    for token in tokens:
        if token.startswith('***') and token.endswith('***'):
            run = p.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith('**') and token.endswith('**'):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = p.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = p.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
        else:
            p.add_run(token)
    return p

def markdown_to_docx(markdown_text: str, docx_path: str):
    """Parse Markdown text and generate a styled .docx document at docx_path."""
    doc = docx.Document()
    
    # Configure page layout (margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = markdown_text.splitlines()
    
    # Parser state variables
    in_table = False
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 1. Handle Table Block
        if stripped.startswith('|'):
            in_table = True
            # Split cells, ignoring outer boundaries
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # Check if this is a separator line (e.g., |---|---|)
            is_separator = all(re.match(r'^:?-+:?$', c) for c in cells) if cells else False
            if not is_separator:
                table_data.append(cells)
            i += 1
            continue
        elif in_table:
            # Table block has ended, flush table to document
            if table_data:
                max_cols = max(len(row) for row in table_data)
                table = doc.add_table(rows=len(table_data), cols=max_cols)
                table.style = 'Table Grid'
                for r_idx, row_cells in enumerate(table_data):
                    for c_idx, val in enumerate(row_cells):
                        if c_idx < max_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            # Add a run inside cell
                            _add_markdown_paragraph(cell, val)
            # Reset table parser state
            in_table = False
            table_data = []
            
        # Ignore empty lines (unless we are flushing tables)
        if not stripped:
            i += 1
            continue

        # 2. Handle Headings
        if stripped.startswith('#'):
            h_match = re.match(r'^(#+)\s+(.*)', stripped)
            if h_match:
                level = len(h_match.group(1))
                title_text = h_match.group(2)
                # Word headings levels are capped at 9, let's limit safely to 6
                safe_level = min(level, 6)
                doc.add_heading(title_text, level=safe_level)
            i += 1
            continue

        # 3. Handle Bullet List
        if stripped.startswith(('- ', '* ')):
            clean_line = stripped[2:]
            _add_markdown_paragraph(doc, clean_line, style='List Bullet')
            i += 1
            continue

        # 4. Handle Numbered List
        list_num_match = re.match(r'^\d+\.\s+(.*)', stripped)
        if list_num_match:
            clean_line = list_num_match.group(1)
            _add_markdown_paragraph(doc, clean_line, style='List Number')
            i += 1
            continue

        # 5. Fallback: Standard Paragraph
        _add_markdown_paragraph(doc, stripped)
        i += 1
        
    # Flush remaining table if document ends on a table
    if in_table and table_data:
        max_cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=max_cols)
        table.style = 'Table Grid'
        for r_idx, row_cells in enumerate(table_data):
            for c_idx, val in enumerate(row_cells):
                if c_idx < max_cols:
                    cell = table.rows[r_idx].cells[c_idx]
                    _add_markdown_paragraph(cell, val)

    doc.save(docx_path)
